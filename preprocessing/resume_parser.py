# Resume Parser Module
# Extracts text from PDF and DOCX resume files with improved parsing.

# PURPOSE: Extract text from resume files (PDF/DOCX) with multiple fallback methods
# COMPONENTS: pypdf, pdfplumber, PyMuPDF, python-docx
# INPUTS: File bytes or file path (PDF/DOCX)
# OUTPUTS: Extracted text string
# WORKFLOW: Try pypdf → pdfplumber → PyMuPDF (fallback chain)
# LOGIC: Multiple extraction methods to handle various PDF/DOCX formats

# WHY: Need reliable text extraction from various resume formats
# WHAT: Converts PDF/DOCX to plain text
# HOW: Uses multiple libraries as fallbacks for maximum compatibility

# Import library for PDF reading
from pypdf import PdfReader
# Import library for DOCX document processing
from docx import Document
# Import library for in-memory file operations
import io
# Import library for regular expressions
import re
# Import type hints
from typing import Optional


def extract_pdf_text_advanced(file_bytes_or_path):
    """
    Extract text from PDF using multiple methods for better accuracy.
    
    Args:
        file_bytes_or_path: File bytes (BytesIO) or file path
    
    Returns:
        str: Extracted text
    """
    text_parts = []
    
    try:
        # Method 1: Try pypdf with improved extraction
        if isinstance(file_bytes_or_path, bytes) or hasattr(file_bytes_or_path, 'read'):
            if hasattr(file_bytes_or_path, 'read'):
                file_bytes = file_bytes_or_path.read()
                file_bytes_or_path.seek(0)  # Reset for next read
            else:
                file_bytes = file_bytes_or_path
            reader = PdfReader(io.BytesIO(file_bytes))
        else:
            reader = PdfReader(file_bytes_or_path)
        
        for page_num, page in enumerate(reader.pages):
            # Try standard extraction
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
            
            # Try extracting with layout preservation (if supported)
            try:
                # Try alternative extraction method with layout mode
                if hasattr(page, 'extract_text'):
                    # Some PDF readers support extraction_mode parameter
                    try:
                        page_text_alt = page.extract_text(extraction_mode="layout")
                        if page_text_alt and len(page_text_alt) > len(page_text):
                            text_parts[-1] = page_text_alt  # Replace with better extraction
                    except TypeError:
                        # extraction_mode not supported, try without
                        pass
            except:
                pass
        
        # Try pdfplumber as fallback if available
        if not text_parts or len(''.join(text_parts)) < 50:
            try:
                import pdfplumber
                if isinstance(file_bytes_or_path, bytes) or hasattr(file_bytes_or_path, 'read'):
                    if hasattr(file_bytes_or_path, 'read'):
                        file_bytes = file_bytes_or_path.read()
                    else:
                        file_bytes = file_bytes_or_path
                    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                        text_parts = [page.extract_text() or "" for page in pdf.pages]
                else:
                    with pdfplumber.open(file_bytes_or_path) as pdf:
                        text_parts = [page.extract_text() or "" for page in pdf.pages]
            except ImportError:
                pass
            except Exception:
                pass
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise Exception(f"PDF extraction error: {str(e)}")


def extract_docx_text_advanced(file_bytes_or_path):
    """
    Extract text from DOCX with improved parsing including tables.
    
    Args:
        file_bytes_or_path: File bytes (BytesIO) or file path
    
    Returns:
        str: Extracted text
    """
    text_parts = []
    
    try:
        # Open document
        if isinstance(file_bytes_or_path, bytes) or hasattr(file_bytes_or_path, 'read'):
            if hasattr(file_bytes_or_path, 'read'):
                file_bytes = file_bytes_or_path.read()
                file_bytes_or_path.seek(0)
            else:
                file_bytes = file_bytes_or_path
            doc = Document(io.BytesIO(file_bytes))
        else:
            doc = Document(file_bytes_or_path)
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            para_text = paragraph.text.strip()
            if para_text:
                text_parts.append(para_text)
        
        # Extract tables
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(" | ".join(row_text))
            if table_text:
                text_parts.append("\n".join(table_text))
        
        # Extract headers and footers
        try:
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        header_text = paragraph.text.strip()
                        if header_text:
                            text_parts.append(header_text)
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        footer_text = paragraph.text.strip()
                        if footer_text:
                            text_parts.append(footer_text)
        except:
            pass
        
        return "\n".join(text_parts)
    
    except Exception as e:
        raise Exception(f"DOCX extraction error: {str(e)}")


def clean_extracted_text(text: str) -> str:
    """
    Clean extracted text while preserving important information.
    
    Args:
        text: Raw extracted text
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove null bytes and control characters (except newlines and tabs)
    text = re.sub(r'[\x00-\x08\x0B-\x0C\x0E-\x1F]', '', text)
    
    # Normalize whitespace but preserve structure
    # Replace multiple spaces with single space
    text = re.sub(r' +', ' ', text)
    
    # Normalize line breaks (keep meaningful breaks)
    text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Multiple newlines to double
    text = re.sub(r'\r\n', '\n', text)  # Windows line endings
    text = re.sub(r'\r', '\n', text)  # Old Mac line endings
    
    # Remove excessive dashes and separators
    text = re.sub(r'-{3,}', '---', text)
    
    # Fix common encoding issues
    text = text.replace('\u2019', "'")  # Right single quotation mark
    text = text.replace('\u2018', "'")  # Left single quotation mark
    text = text.replace('\u201C', '"')  # Left double quotation mark
    text = text.replace('\u201D', '"')  # Right double quotation mark
    text = text.replace('\u2013', '-')  # En dash
    text = text.replace('\u2014', '--')  # Em dash
    text = text.replace('\u2026', '...')  # Ellipsis
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove empty lines at start and end
    text = text.strip()
    
    return text


def extract_resume_text(path_or_file):
    """
    Extract text from resume file (PDF or DOCX) with improved parsing.
    
    Args:
        path_or_file: File path (string) or file-like object (BytesIO)
    
    Returns:
        str: Extracted text from resume
    """
    try:
        # Handle file upload object (Streamlit)
        if hasattr(path_or_file, 'read'):
            file_bytes = path_or_file.read()
            file_bytes_io = io.BytesIO(file_bytes)
            file_extension = path_or_file.name.lower().split('.')[-1]
            
            if file_extension == 'pdf':
                raw_text = extract_pdf_text_advanced(file_bytes_io)
            elif file_extension in ['docx', 'doc']:
                raw_text = extract_docx_text_advanced(file_bytes_io)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            is_pdf = file_extension == 'pdf'
        
        # Handle file path (string)
        else:
            path = str(path_or_file)
            is_pdf = path.lower().endswith('.pdf')
            if is_pdf:
                raw_text = extract_pdf_text_advanced(path)
            elif path.lower().endswith(('.docx', '.doc')):
                raw_text = extract_docx_text_advanced(path)
            else:
                raise ValueError(f"Unsupported file format: {path}")
            file_bytes = None
        
        # Clean and return
        cleaned_text = clean_extracted_text(raw_text)
        
        # OCR fallback for image-based PDFs
        if (not cleaned_text or len(cleaned_text.strip()) < 20) and is_pdf:
            from preprocessing.ocr_parser import try_ocr_on_pdf
            if file_bytes is not None:
                ocr_text = try_ocr_on_pdf(file_bytes)
            else:
                ocr_text = try_ocr_on_pdf(path)
            if ocr_text:
                cleaned_text = clean_extracted_text(ocr_text)
        
        # Validate extraction
        if not cleaned_text or len(cleaned_text.strip()) < 10:
            raise Exception("Extracted text is too short or empty. The file might be corrupted, password-protected, or contain only images.")
        
        return cleaned_text
    
    except ValueError:
        raise
    except Exception as e:
        raise Exception(f"Error extracting resume text: {str(e)}")

